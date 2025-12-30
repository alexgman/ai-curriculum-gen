"""Report Generator - Convert research findings to DOCX and PDF."""
import io
import markdown
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF generation is optional - requires system libraries
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("⚠️ WeasyPrint not available - PDF generation disabled. Install system dependencies for PDF support.")


class ReportGenerator:
    """Generate downloadable reports from markdown research findings."""
    
    def __init__(self):
        self.font_config = FontConfiguration() if WEASYPRINT_AVAILABLE else None
    
    def markdown_to_docx(self, markdown_content: str, title: str = "Research Report") -> io.BytesIO:
        """
        Convert markdown research report to DOCX format.
        
        Args:
            markdown_content: Markdown formatted research report
            title: Report title
        
        Returns:
            BytesIO containing DOCX file
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parse markdown line by line
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add paragraph break
                doc.add_paragraph()
                continue
            
            # Handle headings
            if line.startswith('# '):
                # H1 - Main title
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                # H2 - Section
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # H3 - Subsection
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                # H4
                doc.add_heading(line[5:], level=4)
            
            # Handle lists
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                # Numbered list
                text = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(text, style='List Number')
            
            # Handle bold/italic (simple approach)
            elif '**' in line or '*' in line:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic) to paragraph."""
        # Simple regex-based approach for **bold** and *italic*
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            if i < len(text) - 1 and text[i:i+2] == '**':
                # Bold
                if current:
                    parts.append(("normal", current))
                    current = ""
                # Find closing **
                end = text.find('**', i+2)
                if end != -1:
                    parts.append(("bold", text[i+2:end]))
                    i = end + 2
                    continue
            elif text[i] == '*':
                # Italic
                if current:
                    parts.append(("normal", current))
                    current = ""
                # Find closing *
                end = text.find('*', i+1)
                if end != -1:
                    parts.append(("italic", text[i+1:end]))
                    i = end + 1
                    continue
            
            current += text[i]
            i += 1
        
        if current:
            parts.append(("normal", current))
        
        # Add runs to paragraph
        for style, content in parts:
            run = paragraph.add_run(content)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
    
    def markdown_to_pdf(self, markdown_content: str, title: str = "Research Report") -> io.BytesIO:
        """
        Convert markdown research report to PDF format.
        
        Args:
            markdown_content: Markdown formatted research report
            title: Report title
        
        Returns:
            BytesIO containing PDF file
        
        Raises:
            RuntimeError: If WeasyPrint is not available
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("PDF generation not available. Install system dependencies: sudo apt-get install libpango-1.0-0 libpangocairo-1.0-0")
        # Convert markdown to HTML
        html_content = markdown.markdown(
            markdown_content,
            extensions=['extra', 'nl2br', 'sane_lists']
        )
        
        # Wrap in HTML structure with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                @page {{
                    margin: 1in;
                    size: letter;
                }}
                
                body {{
                    font-family: 'Arial', 'Helvetica', sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                    max-width: 100%;
                }}
                
                h1 {{
                    color: #1a1a1a;
                    font-size: 24pt;
                    margin-top: 0;
                    margin-bottom: 20pt;
                    border-bottom: 2px solid #4A90E2;
                    padding-bottom: 10pt;
                }}
                
                h2 {{
                    color: #2c2c2c;
                    font-size: 18pt;
                    margin-top: 20pt;
                    margin-bottom: 12pt;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 6pt;
                }}
                
                h3 {{
                    color: #3c3c3c;
                    font-size: 14pt;
                    margin-top: 16pt;
                    margin-bottom: 10pt;
                }}
                
                h4 {{
                    color: #4c4c4c;
                    font-size: 12pt;
                    margin-top: 14pt;
                    margin-bottom: 8pt;
                }}
                
                p {{
                    margin-top: 0;
                    margin-bottom: 10pt;
                    text-align: justify;
                }}
                
                ul, ol {{
                    margin-top: 8pt;
                    margin-bottom: 12pt;
                    padding-left: 30pt;
                }}
                
                li {{
                    margin-bottom: 6pt;
                }}
                
                strong {{
                    font-weight: bold;
                    color: #1a1a1a;
                }}
                
                em {{
                    font-style: italic;
                }}
                
                code {{
                    background-color: #f4f4f4;
                    padding: 2pt 4pt;
                    border-radius: 3pt;
                    font-family: 'Courier New', monospace;
                    font-size: 10pt;
                }}
                
                pre {{
                    background-color: #f4f4f4;
                    padding: 12pt;
                    border-radius: 4pt;
                    overflow-x: auto;
                    margin: 12pt 0;
                }}
                
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 12pt 0;
                }}
                
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8pt;
                    text-align: left;
                }}
                
                th {{
                    background-color: #4A90E2;
                    color: white;
                    font-weight: bold;
                }}
                
                tr:nth-child(even) {{
                    background-color: #f9f9f9;
                }}
                
                blockquote {{
                    border-left: 4px solid #4A90E2;
                    padding-left: 16pt;
                    margin-left: 0;
                    font-style: italic;
                    color: #555;
                }}
                
                a {{
                    color: #4A90E2;
                    text-decoration: none;
                }}
                
                a:hover {{
                    text-decoration: underline;
                }}
                
                .page-break {{
                    page-break-after: always;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        pdf_buffer = io.BytesIO()
        HTML(string=full_html).write_pdf(
            pdf_buffer,
            font_config=self.font_config
        )
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    def generate_filename(self, topic: str, format: str = "pdf") -> str:
        """
        Generate a clean filename for the report.
        
        Args:
            topic: Research topic
            format: File format (pdf or docx)
        
        Returns:
            Filename string
        """
        # Clean topic string
        clean_topic = re.sub(r'[^\w\s-]', '', topic).strip()
        clean_topic = re.sub(r'[-\s]+', '_', clean_topic)
        
        # Limit length
        if len(clean_topic) > 50:
            clean_topic = clean_topic[:50]
        
        return f"curriculum_research_{clean_topic}.{format}"


# Singleton instance
report_generator = ReportGenerator()

